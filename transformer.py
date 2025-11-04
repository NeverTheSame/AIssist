import csv
import json
import sys
from datetime import datetime
import os
import logging
from bs4 import BeautifulSoup
from docx import Document
from openai import AzureOpenAI
from config import config

# Configure logging
def setup_logging():
    """Setup logging for incident dumper"""
    # Create logs directory if it doesn't exist
    os.makedirs("logs", exist_ok=True)
    
    # Create logger
    logger = logging.getLogger(__name__)
    logger.setLevel(logging.INFO)
    
    # Clear any existing handlers
    logger.handlers.clear()
    
    # Create formatter
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    
    # File handler - detailed logging
    file_handler = logging.FileHandler('logs/transformer.log', encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    file_handler.setFormatter(formatter)
    logger.addHandler(file_handler)
    
    # Console handler - only warnings and errors
    console_handler = logging.StreamHandler(sys.stdout)
    console_handler.setLevel(logging.WARNING)
    console_handler.setFormatter(logging.Formatter('%(levelname)s: %(message)s'))
    logger.addHandler(console_handler)
    
    return logger

logger = setup_logging()

def clean_html_content(html_content):
    """Clean HTML content by removing encrypted images and rich text formatting"""
    if not html_content:
        return ""
    
    import re
    
    # First, use regex to remove any image tags with data:image/png or REDACTED content
    # This handles both properly closed and malformed HTML where tags might not be properly closed
    text = re.sub(r'<img[^>]*src=["\']?data:image/png[^>]*>', '', html_content, flags=re.IGNORECASE)
    text = re.sub(r'<img[^>]*src=["\']?[^"\']*REDACTED[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<img[^>]*src=["\']?[^"\']*data:image/png[^>]*>', '', text, flags=re.IGNORECASE)
    
    # Handle malformed img tags that don't have proper closing > (like at end of lines)
    text = re.sub(r'<img[^>]*src=["\']?data:image/png[^>]*[^>]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<img[^>]*src=["\']?[^"\']*REDACTED[^>]*[^>]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<img[^>]*src=["\']?[^"\']*data:image/png[^>]*[^>]', '', text, flags=re.IGNORECASE)
    
    # Also remove any remaining img tags that might be very long (likely base64)
    text = re.sub(r'<img[^>]*src=["\']?[^"\']{1000,}[^>]*>', '', text, flags=re.IGNORECASE)
    text = re.sub(r'<img[^>]*src=["\']?[^"\']{1000,}[^>]*[^>]', '', text, flags=re.IGNORECASE)
    
    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(text, "html.parser")
    
    # Remove all remaining img tags with data:image/png src (encrypted images)
    for img in soup.find_all('img'):
        src = img.get('src', '')
        # Check for various patterns of encrypted/encoded images
        if (src.startswith('data:image/png') or 
            'data:image/png;base64' in src or 
            'REDACTED' in src or
            len(src) > 1000):  # Very long src attributes are likely base64 images
            logger.info(f"Removing encrypted/encoded image: {src[:50]}...")
            img.decompose()
    
    # Remove other potentially problematic elements
    for element in soup.find_all(['script', 'style', 'link']):
        element.decompose()
    
    # Convert to plain text, preserving some structure
    # Replace common HTML elements with appropriate text formatting
    for tag in soup.find_all(['p', 'div', 'br']):
        if tag.name == 'br':
            tag.replace_with('\n')
        elif tag.name in ['p', 'div']:
            # Add newlines around paragraphs and divs
            tag.insert_before('\n')
            tag.insert_after('\n')
    
    # Get text and clean up whitespace
    text = soup.get_text(separator=' ', strip=True)
    
    # Clean up excessive whitespace and newlines
    text = re.sub(r'\n\s*\n', '\n\n', text)  # Replace multiple newlines with double newlines
    text = re.sub(r' +', ' ', text)  # Replace multiple spaces with single space
    text = text.strip()
    
    return text

def parse_multi_section_csv(file_path):
    sections = {}
    current_section = None
    current_headers = None
    
    # Increase CSV field size limit to handle large fields
    import csv
    maxInt = sys.maxsize
    while True:
        try:
            csv.field_size_limit(maxInt)
            break
        except OverflowError:
            maxInt = int(maxInt/10)
    
    try:
        with open(file_path, newline='', encoding='utf-8') as csvfile:
            reader = csv.reader(csvfile)
            for row_num, row in enumerate(reader, 1):
                try:
                    if not row or all(cell.strip() == '' for cell in row):
                        continue  # skip blank lines
                    if row[0].startswith('---') and row[0].endswith('---'):
                        # New section header
                        current_section = row[0].strip('- ').strip()
                        sections[current_section] = []
                        current_headers = None
                    elif current_section and current_headers is None:
                        current_headers = row
                    elif current_section and current_headers:
                        # Truncate any fields that are too long
                        processed_row = []
                        for cell in row:
                            if isinstance(cell, str) and len(cell) > 100000:  # Limit to 100KB per field
                                logger.warning(f"Truncating field at row {row_num} that was {len(cell)} characters long")
                                cell = cell[:100000] + "... [TRUNCATED]"
                            processed_row.append(cell)
                        
                        item = {header: value for header, value in zip(current_headers, processed_row)}
                        sections[current_section].append(item)
                except Exception as e:
                    logger.warning(f"Error processing row {row_num}: {e}. Skipping this row.")
                    continue
    except Exception as e:
        logger.error(f"Error reading CSV file {file_path}: {e}")
        raise
    
    return sections

def dump_discussion_items_to_json(discussion_items, incident_number, output_dir="processed_incidents", summary_content=None):
    """Dump incidents to a JSON file in a clean format, including internal AI summary if provided. (No summary field in output)"""
    # Create output directory if it doesn't exist
    os.makedirs(output_dir, exist_ok=True)
    
    # Clean summary content if provided
    cleaned_summary = None
    if summary_content:
        logger.info("Cleaning summary content...")
        cleaned_summary = clean_html_content(summary_content)
        logger.info(f"Summary content cleaned. Original length: {len(summary_content)}, Cleaned length: {len(cleaned_summary)}")
    
    # Format the data for AI processing
    formatted_discussion_items = []
    filtered_count = 0
    for discussion_item in discussion_items:
        raw_content = (discussion_item.get('Text') or discussion_item.get('content') or '').strip()
        # Clean HTML/rich text using BeautifulSoup
        if raw_content:
            clean_content = clean_html_content(raw_content)
        else:
            clean_content = ''
        
        # Get author information
        author = discussion_item.get('ChangedBy') or discussion_item.get('author')
        
        # Filter out unwanted entries
        # Skip entries with author "gautosvc" and content starting with "Support ICM enrichment CEM MDE"
        if author == "gautosvc" and clean_content.startswith("Support ICM enrichment CEM MDE"):
            logger.info(f"Filtering out gautosvc entry with Support ICM enrichment content")
            filtered_count += 1
            continue
        
        formatted_discussion_item = {
            "timestamp": discussion_item.get('Date') or discussion_item.get('timestamp'),
            "author": author,
            "content": clean_content
        }
        formatted_discussion_items.append(formatted_discussion_item)
    
    # Sort by timestamp
    formatted_discussion_items.sort(key=lambda x: x['timestamp'])
    

    
    # Create the output file (without 'incident_' prefix)
    output_file = os.path.join(output_dir, f"{incident_number}.json")
    
    # Write to file (no summary field)
    with open(output_file, 'w', encoding='utf-8') as f:
        output_data = {
            "incident_id": incident_number,
            "total_entries": len(formatted_discussion_items),
            "conversation": formatted_discussion_items
        }
        
        # Add cleaned summary if available
        if cleaned_summary:
            output_data["summary"] = cleaned_summary
            logger.info(f"Added cleaned summary to output (length: {len(cleaned_summary)})")
            print(f"[transformer] Final summary saved ({len(cleaned_summary)} chars)")

        
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    
    logger.info(f"Processed {len(formatted_discussion_items)} entries (filtered out {filtered_count} unwanted entries)")
    print(f"✅ Created: {output_file}")
    
    return output_file, filtered_count

def is_summary_redacted(summary_content):
    """
    Check if the summary content is redacted.
    Returns True if summary contains the redaction pattern.
    The user mentioned it comes back as: "Summary": ** REDACTED **
    """
    if not summary_content or not isinstance(summary_content, str):
        return False
    
    summary_stripped = summary_content.strip()

    # Only treat as redacted when it exactly matches the token "** REDACTED **" (ignoring outer whitespace)
    # Allow flexible spacing around the word between the asterisks
    import re
    pattern = r"^\s*\*\*\s*REDACTED\s*\*\*\s*$"
    if re.match(pattern, summary_stripped, flags=re.IGNORECASE):
        logger.info("Detected exact redacted token '** REDACTED **'")
        return True

    return False

def prompt_user_for_docx(incident_number):
    """
    Return the fixed manual.docx path from project root if it exists.
    No interactive prompts are used.
    """
    # Use project root from config and fixed filename manual.docx
    fixed_docx_path = os.path.join(str(config.root_dir), "manual.docx")
    
    if os.path.exists(fixed_docx_path):
        print(f"Using manual.docx from project root: {fixed_docx_path}")
        return fixed_docx_path
    
    print("manual.docx not found at project root. Continuing with redacted summary.")
    return None

def extract_text_from_docx(docx_path):
    """
    Extract text content from manual file. Prefer DOCX parsing; if that fails,
    fall back to reading the file as UTF-8 text.
    Returns the text content as a string.
    """
    # First try DOCX parsing
    try:
        logger.info(f"Extracting text from docx: {docx_path}")
        doc = Document(docx_path)
        # Extract text from all paragraphs
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text and paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        # Also extract text from tables if any
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        text_parts.append(cell.text.strip())
        full_text = "\n".join(text_parts)
        logger.info(f"Extracted {len(full_text)} characters from docx")
        if not full_text.strip():
            logger.warning("Docx file appears to be empty or contains no extractable text")
        return full_text
    except Exception as e:
        logger.warning(f"DOCX parse failed, attempting plain-text fallback: {e}")
        # Fallback: treat the file as plain text (useful if user saved .docx as plain txt)
        try:
            with open(docx_path, 'r', encoding='utf-8') as tf:
                txt = tf.read()
            logger.info(f"Extracted {len(txt)} characters from plain-text fallback")
            return txt
        except Exception as e2:
            logger.error(f"Error reading manual file as text: {e2}")
            raise Exception(f"Failed to extract text from manual file: {e2}")

def extract_summary_from_docx_text(docx_text):
    """
    Use Azure Router (GPT-5) to extract the Authored Summary from docx text.
    The LLM will identify and extract the summary content.
    """
    try:
        logger.info("Initializing Azure Router client for summary extraction")
        
        # Initialize Azure OpenAI client (Azure Router)
        client = AzureOpenAI(
            api_key=config.ai_service_api_key,
            api_version=config.ai_service_api_version,
            azure_endpoint=config.ai_service_endpoint
        )
        
        # Create prompt for LLM to extract ALL content from manual docx (verbatim content extraction)
        system_prompt = (
            "You are a document processing assistant. Your task is to extract the ENTIRE textual content "
            "from the provided document without omitting any relevant information. Preserve order and main "
            "paragraph breaks. Do not add analysis or commentary."
        )

        user_prompt = (
            "Extract all textual content from the following document text. Return only the content, no labels or extra text.\n\n"
            f"{docx_text}"
        )

        # If no text is available from manual file, do not call LLM
        if not docx_text or not docx_text.strip():
            logger.warning("Manual file is empty after extraction. Skipping LLM extraction.")
            print("⚠️  manual.docx appears empty. Populate it with the Authored Summary and retry.")
            return None

        logger.info("Calling Azure Router to extract summary from docx text")
        print("Processing document with LLM to extract all manual.docx content...")
        # Show a preview of the LLM request (system + user prompt)
        sys_preview = system_prompt[:600].replace('\n', ' ').replace('\r', ' ')
        usr_preview = user_prompt[:1000].replace('\n', ' ').replace('\r', ' ')
        print(f"[transformer] LLM system prompt preview: {sys_preview}")
        print(f"[transformer] LLM user prompt preview: {usr_preview}")
        
        # Call Azure Router
        # For Azure OpenAI, pass the deployment name via the 'model' parameter
        response = client.chat.completions.create(
            model=config.ai_service_deployment_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
            max_tokens=4000
        )
        
        extracted_summary = response.choices[0].message.content.strip()
        
        logger.info(f"Successfully extracted summary from docx (length: {len(extracted_summary)})")
        # Print a visible preview of what LLM produced
        preview = extracted_summary[:800].replace('\n', ' ').replace('\r', ' ')
        print(f"[transformer] LLM output preview: {preview}")
        
        if "No Authored Summary found" in extracted_summary:
            logger.warning("LLM could not find Authored Summary in document")
            print("⚠️  LLM could not find Authored Summary in the document.")
            return None
        
        print(f"✅ Successfully extracted Authored Summary ({len(extracted_summary)} characters)")
        return extracted_summary
    
    except Exception as e:
        logger.error(f"Error extracting summary using LLM: {e}")
        print(f"❌ Error processing document with LLM: {e}")
        raise Exception(f"Failed to extract summary from docx using LLM: {e}")

def main():
    import argparse
    
    parser = argparse.ArgumentParser(description='Process and dump incident data to JSON format.')
    parser.add_argument('file_path', help='Path to the CSV file (e.g., icms/12345/12345.csv or icms/12345.csv)')
    parser.add_argument('--output-dir', default='processed_incidents', 
                      help='Directory to store processed incident files')
    args = parser.parse_args()
    
    try:
        # Parse the multi-section CSV
        sections = parse_multi_section_csv(args.file_path)
        
        # Debug: Log all available sections
        logger.info(f"Available sections in CSV: {list(sections.keys())}")
        for section_name, section_data in sections.items():
            logger.info(f"Section '{section_name}': {len(section_data)} items")
        
        # Extract incident number from filename
        filename = os.path.basename(args.file_path)
        incident_number = ''.join(filter(str.isdigit, filename))
        
        # Discussions section
        discussion_items = sections.get('Discussions', [])
        
        # Authored summary section (new field from updated query)
        summary_content = None
        summary_section = sections.get('Authored summary', [])
        if summary_section:
            # Prefer aggregate across all rows/columns to handle multi-line/row summaries
            try:
                aggregated = []
                for row_item in summary_section:
                    # Join all cell values in the row, preserving order of headers when possible
                    if isinstance(row_item, dict):
                        row_text = " ".join([str(v) for v in row_item.values() if v])
                    else:
                        row_text = str(row_item)
                    if row_text.strip():
                        aggregated.append(row_text.strip())
                agg_text = "\n".join(aggregated).strip()
            except Exception as e:
                logger.warning(f"Failed to aggregate authored summary rows: {e}")
                agg_text = ''

            # Backward-compat single-cell fallback
            single_cell = summary_section[0].get('Summary', '') if isinstance(summary_section[0], dict) else ''

            # Choose the longer non-empty variant
            summary_content = agg_text if len(agg_text) >= len(single_cell) else single_cell
            logger.info(f"Found authored summary content (length: {len(summary_content)})")
        
        # Check if summary is redacted
        if summary_content and is_summary_redacted(summary_content):
            logger.warning(f"Detected redacted summary for incident {incident_number}")
            print(f"\n⚠️  WARNING: Authored summary for incident {incident_number} is REDACTED")
            
            # Prompt user for manual.docx
            docx_path = prompt_user_for_docx(incident_number)
            
            if docx_path:
                try:
                    # Extract text from docx
                    docx_text = extract_text_from_docx(docx_path)
                    logger.info(f"manual.docx path: {docx_path}")
                    logger.info(f"manual.docx extracted text length: {len(docx_text)}")
                    preview_in = docx_text[:500].replace('\n', ' ').replace('\r', ' ')
                    print(f"[transformer] Read manual.docx ({len(docx_text)} chars). Preview: {preview_in}")
                    
                    # Use LLM to extract content from docx text
                    extracted_summary = extract_summary_from_docx_text(docx_text)
                    
                    if extracted_summary:
                        # Replace redacted summary with extracted one
                        logger.info(f"Replacing redacted summary with extracted summary from docx")
                        logger.info(f"Extracted summary length: {len(extracted_summary)}")
                        preview_out = extracted_summary[:800].replace('\n', ' ').replace('\r', ' ')
                        print(f"[transformer] LLM extracted summary ({len(extracted_summary)} chars). Preview: {preview_out}")
                        summary_content = extracted_summary
                        print(f"✅ Authored Summary extracted and will be used instead of redacted content")
                    else:
                        # Fallback: if LLM extraction failed but manual has text, use it directly
                        if docx_text and docx_text.strip():
                            logger.warning("LLM extraction failed; using manual.docx content directly as summary")
                            print("⚠️  LLM extraction failed; using manual.docx content directly as summary.")
                            summary_content = docx_text.strip()
                        else:
                            logger.warning("Could not extract summary from docx, proceeding with redacted summary")
                            print("⚠️  Could not extract summary from docx. Continuing with redacted summary.")
                
                except Exception as e:
                    logger.error(f"Error processing manual.docx: {e}")
                    print(f"❌ Error processing manual.docx: {e}")
                    print("Continuing with redacted summary...")
                    # Continue with redacted summary if processing fails
            else:
                print("⚠️  manual.docx not found at project root. Skipping manual extraction.")
        
        # Dump to JSON
        output_file, filtered_count = dump_discussion_items_to_json(
            discussion_items, 
            incident_number,
            args.output_dir,
            summary_content
        )
        
        print(f"\nProcessed incident data has been saved to: {output_file}")
        print(f"Total entries processed: {len(discussion_items)}")
        if filtered_count > 0:
            print(f"Filtered out {filtered_count} gautosvc entries with Support ICM enrichment content")
        if summary_content:
            print("Authored summary content has been cleaned and included in the JSON file")
        
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    main() 